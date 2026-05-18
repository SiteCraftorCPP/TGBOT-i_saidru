from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate

from app.core.config import Settings
from app.core.doc_filename import build_dynamic_docx_filename
from app.schemas.ai import TemplateMeta
from app.services.catalog import TemplateCatalog
from app.services.deepseek import DeepSeekClient
from app.services.pdf import PdfGenerationError, convert_docx_to_pdf


class DocumentGenerator:
    def __init__(self, settings: Settings, catalog: TemplateCatalog, deepseek: DeepSeekClient):
        self.settings = settings
        self.catalog = catalog
        self.deepseek = deepseek

    async def generate(
        self,
        *,
        document_id: int,
        template: TemplateMeta,
        raw_answers: dict[str, str],
    ) -> tuple[dict[str, Any], str, str, Path, Path | None]:
        fill = await self.deepseek.normalize_answers(template, raw_answers)
        values = {field.key: fill.values.get(field.key, "") for field in template.fields}
        document_text = self.render_text(template, values)

        output_dir = self.settings.storage_dir / str(document_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        docx_path = output_dir / f"{template.document_type}_{document_id}.docx"

        tpl = DocxTemplate(self.catalog.path_for(template))
        tpl.render(values)
        tpl.save(docx_path)

        pdf_path: Path | None = None
        try:
            pdf_path = await convert_docx_to_pdf(docx_path, output_dir, self.settings)
        except (FileNotFoundError, PdfGenerationError):
            # DOCX is still a valid deliverable; PDF can be regenerated after LibreOffice setup.
            pdf_path = None

        instruction = fill.instruction.strip() or template.instruction
        return values, document_text, instruction, docx_path, pdf_path

    async def generate_dynamic(
        self,
        *,
        document_id: int,
        request_text: str,
        details_text: str,
    ) -> tuple[str, str, Path, Path | None]:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        result = await self.deepseek.generate_dynamic_document(request_text, details_text)
        instruction = result.instruction

        output_dir = self.settings.storage_dir / str(document_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        title_for_file = (
            (result.title or "").strip()
            or (result.subtitle or "").strip()
            or "документ"
        )
        docx_path = output_dir / build_dynamic_docx_filename(
            title=title_for_file,
            request_text=request_text,
            details_text=details_text,
        )

        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Set 1.5 line spacing for the whole document
        style.paragraph_format.line_spacing = 1.5

        # Set page margins (ГОСТ Р 7.0.97-2016)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(1.5)

        document_parts = []

        # Header (Шапка)
        # We use a table to place the header strictly on the right half of the page
        if result.header:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Cm(8.5)
            table.columns[1].width = Cm(8.0)
            
            cell = table.cell(0, 1)
            bold_prefixes = [
                "Истец:", "Ответчик:", "Заявитель:", "Жалобщик:", "Истец ",
                "В суд:", "Мировому судье:", "Кому:", "От кого:",
                "Взыскатель:", "Должник:", "Покупатель:", "Продавец:",
                "Арендодатель:", "Арендатор:", "Наймодатель:", "Наниматель:",
                "Займодатель:", "Заёмщик:", "Заемщик:",
                "Заказчик:", "Исполнитель:", "Подрядчик:", "Доверитель:", "Поверенный:",
            ]
            
            for i, line in enumerate(result.header):
                p = cell.add_paragraph() if i > 0 else cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                
                matched = False
                for prefix in bold_prefixes:
                    if line.startswith(prefix):
                        run_bold = p.add_run(prefix)
                        run_bold.bold = True
                        run_normal = p.add_run(line[len(prefix):])
                        matched = True
                        break
                if not matched:
                    p.add_run(line)
                
                document_parts.append(line)
        
        # Title
        if result.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(result.title.upper())
            run.bold = True
            document_parts.append("\n" + result.title.upper())

        # Subtitle
        if result.subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(result.subtitle)
            run.bold = True
            document_parts.append(result.subtitle)

        # Body
        document_parts.append("")
        for para in result.body:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(0)
            document_parts.append(para)

        # Requests
        if result.requests:
            p = doc.add_paragraph("ПРОШУ:")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            document_parts.append("\nПРОШУ:")
            for req in result.requests:
                p = doc.add_paragraph(req)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.space_after = Pt(0)
                document_parts.append(req)

        # Attachments
        if result.attachments:
            p = doc.add_paragraph("Приложения:")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.runs[0].bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            document_parts.append("\nПриложения:")
            for i, att in enumerate(result.attachments, 1):
                # Ensure attachments are numbered if not already
                text = att if att.strip().startswith(str(i)) else f"{i}. {att}"
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.space_after = Pt(0)
                document_parts.append(text)

        # Date and signature
        if result.date_and_signature:
            # Create a table for date and signature to spread them out nicely
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Cm(8.25)
            table.columns[1].width = Cm(8.25)
            
            # Add some space before the table
            p_before = doc.add_paragraph()
            p_before.paragraph_format.space_before = Pt(24)
            
            # Try to split date and signature if they are separated by spaces
            parts = result.date_and_signature.split("    ")
            if len(parts) >= 2:
                date_part = parts[0].strip()
                sig_part = "    ".join(parts[1:]).strip()
            else:
                date_part = result.date_and_signature
                sig_part = ""

            cell_left = table.cell(0, 0)
            p_left = cell_left.paragraphs[0]
            p_left.add_run(date_part)
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            cell_right = table.cell(0, 1)
            p_right = cell_right.paragraphs[0]
            p_right.add_run(sig_part)
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            document_parts.append("\n" + result.date_and_signature)
        
        doc.save(docx_path)
        
        document_text = "\n".join(document_parts)

        pdf_path: Path | None = None
        try:
            pdf_path = await convert_docx_to_pdf(docx_path, output_dir, self.settings)
        except (FileNotFoundError, PdfGenerationError):
            pdf_path = None

        return document_text, instruction, docx_path, pdf_path

    @staticmethod
    def render_text(template: TemplateMeta, values: dict[str, Any]) -> str:
        text = "\n\n".join(template.body)
        for key, value in values.items():
            text = text.replace("{{ " + key + " }}", str(value))
            text = text.replace("{{" + key + "}}", str(value))
        return text
